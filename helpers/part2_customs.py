# docxtpl is limited in its ability to change complex tables.  This is a solution directly using python-docx
import docx.document
from docx import Document

import report_input
from helpers.part_2.table12_14 import get_table12_14
from helpers.part_2.table13 import place_table13

context_functions = {
    "table13": place_table13,
}

def apply_custom_tables(template_name: str) -> None:
    doc: docx.document.Document = Document(template_name)
    for table in doc.tables:
        for keys in context_functions.keys():
            if f"<<{keys}>>" in table.cell(0, 0).text:
                # if it contains the tag we are looking for, remove it
                for paragraphs in table.cell(0, 0).paragraphs:
                    for run in paragraphs.runs:
                        if f"<<{keys}>>" in run.text:
                            run.text = run.text.replace(f"<<{keys}>>", "")
                # then apply the custom table
                context_functions[keys](table)



    doc.save("custom_tables.docx")
    return None
